"""Document generation for research tasks."""
import os
from io import BytesIO
from pathlib import Path


def _safe(text: str) -> str:
    """Replace characters outside Latin-1 with '?' so reportlab's default
    Helvetica font doesn't throw UnicodeEncodeError."""
    return text.encode('latin-1', errors='replace').decode('latin-1')


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip('#')
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _strip_leading_bullets(text: str) -> str:
    """Remove leading bullet/dash/special chars that aren't printable ASCII letters or digits."""
    import re
    return re.sub(r'^[\s\W•\-–—*#>]+', '', text).strip()


def _build_summary_section(sections: list[dict], summary_as_bullets: bool) -> dict:
    """Build a summary section from the first ~100 chars of each section's content."""
    bullets = []
    for sec in sections:
        content = sec.get('content', '')
        # Take first line or first 100 chars, whichever is shorter
        first_line = content.split('\n')[0].strip()
        snippet = first_line[:100] if first_line else content[:100].strip()
        snippet = _strip_leading_bullets(snippet)
        if snippet:
            bullets.append(snippet)
    if summary_as_bullets:
        summary_content = '\n'.join(f"• {b}" for b in bullets)
    else:
        summary_content = '\n'.join(bullets)
    return {"heading": "Summary", "content": summary_content}


def _render_pdf_section_content(content: str, section_style: str, body_style) -> list:
    """Return a list of Paragraph flowables based on section_style."""
    from reportlab.platypus import Paragraph

    lines = [l.strip() for l in content.split('\n') if l.strip()]
    result = []
    if section_style == "bullets":
        for line in lines:
            result.append(Paragraph(f"• {line}", body_style))
    elif section_style == "mixed":
        for i, line in enumerate(lines):
            if i == 0:
                result.append(Paragraph(line, body_style))
            else:
                result.append(Paragraph(f"• {line}", body_style))
    else:  # "paragraphs" (default)
        for line in lines:
            result.append(Paragraph(line, body_style))
    return result


def _fetch_image_bytes(url: str) -> BytesIO | None:
    """Download an image URL and return its bytes, or None on failure."""
    try:
        import httpx
        resp = httpx.get(url, follow_redirects=True, timeout=8)
        if resp.status_code == 200 and resp.headers.get("content-type", "").startswith("image/"):
            return BytesIO(resp.content)
    except Exception:
        pass
    return None


def generate_pdf(title: str, sections: list[dict], sources: list[str],
                 output_path: str, fmt: dict,
                 image_urls: list[str] | None = None) -> None:
    """Generate a PDF report using reportlab."""
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Image as RLImage
    from reportlab.lib.colors import HexColor

    pagesize = A4 if fmt.get("page_size") == "a4" else letter

    doc = SimpleDocTemplate(output_path, pagesize=pagesize,
                            rightMargin=inch, leftMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    styles = getSampleStyleSheet()

    # Custom styles driven by fmt
    title_style = ParagraphStyle('CustomTitle', parent=styles['Title'],
                                  fontSize=fmt["title_font_size"], spaceAfter=20)
    heading_style = ParagraphStyle('CustomHeading', parent=styles['Heading2'],
                                    fontSize=fmt["heading_font_size"], spaceBefore=16, spaceAfter=8,
                                    textColor=HexColor(fmt["header_color"]))
    body_font_size = fmt["body_font_size"]
    body_style = ParagraphStyle('CustomBody', parent=styles['Normal'],
                                 fontSize=body_font_size,
                                 leading=round(body_font_size * 1.45),
                                 spaceAfter=8)
    source_style = ParagraphStyle('Source', parent=styles['Normal'],
                                   fontSize=9, textColor=HexColor('#7f8c8d'))

    story = []
    story.append(Paragraph(_safe(title), title_style))
    story.append(HRFlowable(width="100%", thickness=2, color=HexColor(fmt["accent_color"])))
    story.append(Spacer(1, 0.2 * inch))

    # Build sections list, optionally prepending a summary
    all_sections = list(sections)
    if fmt.get("include_summary") and all_sections:
        summary_sec = _build_summary_section(all_sections, fmt.get("summary_as_bullets", True))
        all_sections = [summary_sec] + all_sections

    section_style_val = fmt.get("section_style", "paragraphs")

    for section in all_sections:
        story.append(Paragraph(_safe(section['heading']), heading_style))
        story.extend(_render_pdf_section_content(_safe(section['content']), section_style_val, body_style))
        story.append(Spacer(1, 0.1 * inch))

    # Embed images if requested
    if image_urls:
        embedded = 0
        for url in image_urls:
            if embedded >= 5:
                break
            img_bytes = _fetch_image_bytes(url)
            if img_bytes is None:
                continue
            try:
                max_width = 4.5 * inch
                img = RLImage(img_bytes, width=max_width)
                # Preserve aspect ratio: compute height from reportlab's detected size
                img_w, img_h = img.imageWidth, img.imageHeight
                if img_w and img_h:
                    img.drawHeight = max_width * img_h / img_w
                    img.drawWidth = max_width
                story.append(Spacer(1, 0.15 * inch))
                story.append(img)
                embedded += 1
            except Exception:
                pass

    if fmt.get("include_links") and sources:
        story.append(HRFlowable(width="100%", thickness=1, color=HexColor('#bdc3c7')))
        story.append(Spacer(1, 0.1 * inch))
        story.append(Paragraph("Sources", heading_style))
        for i, src in enumerate(sources, 1):
            story.append(Paragraph(f"{i}. {_safe(src)}", source_style))

    doc.build(story)


def generate_docx(title: str, sections: list[dict], sources: list[str],
                  output_path: str, fmt: dict,
                  image_urls: list[str] | None = None) -> None:
    """Generate a Word document using python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Page size
    if fmt.get("page_size") == "a4":
        section_obj = doc.sections[0]
        section_obj.page_width = Mm(210)
        section_obj.page_height = Mm(297)

    # Title
    title_para = doc.add_heading(_safe(title), 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(fmt["title_font_size"])

    doc.add_paragraph()  # spacer

    # Parse heading color once
    header_rgb = _hex_to_rgb(fmt["header_color"])

    # Build sections list, optionally prepending a summary
    all_sections = list(sections)
    if fmt.get("include_summary") and all_sections:
        summary_sec = _build_summary_section(all_sections, fmt.get("summary_as_bullets", True))
        all_sections = [summary_sec] + all_sections

    section_style_val = fmt.get("section_style", "paragraphs")
    body_pt = Pt(fmt["body_font_size"])

    for section in all_sections:
        heading_para = doc.add_heading(_safe(section['heading']), level=1)
        for run in heading_para.runs:
            run.font.size = Pt(fmt["heading_font_size"])
            run.font.color.rgb = RGBColor(*header_rgb)

        lines = [l.strip() for l in _safe(section['content']).split('\n') if l.strip()]
        for i, line in enumerate(lines):
            if section_style_val == "bullets":
                p = doc.add_paragraph(style='List Bullet')
                run = p.add_run(line)
                run.font.size = body_pt
            elif section_style_val == "mixed":
                if i == 0:
                    p = doc.add_paragraph()
                    run = p.add_run(line)
                    run.font.size = body_pt
                else:
                    p = doc.add_paragraph(style='List Bullet')
                    run = p.add_run(line)
                    run.font.size = body_pt
            else:  # paragraphs
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = body_pt

    # Embed images if requested
    if image_urls:
        embedded = 0
        for url in image_urls:
            if embedded >= 5:
                break
            img_bytes = _fetch_image_bytes(url)
            if img_bytes is None:
                continue
            try:
                doc.add_picture(img_bytes, width=Inches(4.5))
                embedded += 1
            except Exception:
                pass

    if fmt.get("include_links") and sources:
        sources_heading = doc.add_heading('Sources', level=1)
        for run in sources_heading.runs:
            run.font.size = Pt(fmt["heading_font_size"])
            run.font.color.rgb = RGBColor(*header_rgb)
        for i, src in enumerate(sources, 1):
            p = doc.add_paragraph(style='List Number')
            p.add_run(f"{src}").font.size = Pt(9)

    doc.save(output_path)


def _title_to_filename(title: str) -> str:
    """Convert a report title to a safe filename (no extension)."""
    import re
    # Keep alphanumerics, spaces, hyphens; replace everything else
    safe = re.sub(r'[^\w\s\-]', '', title)
    # Collapse whitespace to single underscore
    safe = re.sub(r'\s+', '_', safe.strip())
    # Truncate so the full path stays reasonable
    return safe[:80] if safe else "report"


def generate_document(title: str, sections: list[dict], sources: list[str],
                       output_format: str, output_dir: str, fmt: dict,
                       image_urls: list[str] | None = None) -> str:
    """Generate document and return full path."""
    os.makedirs(output_dir, exist_ok=True)
    ext = "pdf" if output_format == "pdf" else "docx"
    filename = f"{_title_to_filename(title)}.{ext}"
    full_path = os.path.join(output_dir, filename)

    if output_format == "pdf":
        generate_pdf(title, sections, sources, full_path, fmt, image_urls=image_urls)
    else:
        generate_docx(title, sections, sources, full_path, fmt, image_urls=image_urls)

    return full_path
