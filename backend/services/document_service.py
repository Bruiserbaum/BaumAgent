"""Extract text content from uploaded documents (PDF, DOCX, XLSX, CSV)."""
import csv
import io


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".xls", ".csv"}
SUPPORTED_MIMETYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
    "text/csv",
    "application/csv",
}


def _extension(filename: str) -> str:
    """Return lowercase file extension including the dot."""
    idx = filename.rfind(".")
    if idx == -1:
        return ""
    return filename[idx:].lower()


def extract_text(filename: str, data: bytes) -> str:
    """
    Extract human-readable text from a document.

    Raises ValueError if the file type is not supported.
    """
    ext = _extension(filename)

    if ext == ".pdf":
        return _extract_pdf(data)
    elif ext == ".docx":
        return _extract_docx(data)
    elif ext in (".xlsx", ".xls"):
        return _extract_xlsx(data)
    elif ext == ".csv":
        return _extract_csv(data)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(data: bytes) -> str:
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ValueError("PyPDF2 is not installed — cannot parse PDF files")

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"--- Page {i + 1} ---\n{text.strip()}")
    if not pages:
        return "(No extractable text found in the PDF)"
    return "\n\n".join(pages)


def _extract_docx(data: bytes) -> str:
    """Extract text from a Word .docx file."""
    try:
        from docx import Document
    except ImportError:
        raise ValueError("python-docx is not installed — cannot parse DOCX files")

    doc = Document(io.BytesIO(data))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            paragraphs.append("\n".join(rows))

    if not paragraphs:
        return "(No extractable text found in the document)"
    return "\n\n".join(paragraphs)


def _extract_xlsx(data: bytes) -> str:
    """Extract text from an Excel .xlsx file."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ValueError("openpyxl is not installed — cannot parse Excel files")

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sheets = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(c for c in cells):
                rows.append(" | ".join(cells))
        if rows:
            sheets.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))
    wb.close()

    if not sheets:
        return "(No data found in the spreadsheet)"
    return "\n\n".join(sheets)


def _extract_csv(data: bytes) -> str:
    """Extract text from a CSV file."""
    # Try to decode as UTF-8, fall back to latin-1
    try:
        text = data.decode("utf-8")
    except UnicodeDecodeError:
        text = data.decode("latin-1")

    reader = csv.reader(io.StringIO(text))
    rows = []
    for row in reader:
        cells = [c.strip() for c in row]
        if any(c for c in cells):
            rows.append(" | ".join(cells))

    if not rows:
        return "(No data found in the CSV file)"
    return "\n".join(rows)
